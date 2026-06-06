"""DOCX renderer (flow fidelity). Rebuilds an editable Word doc from the IR.

Layout = logical structure (headings, paragraphs, lists, tables), not pixel-exact — the
honest trade-off for an *editable* target. For visual fidelity use the PDF overlay renderer.
"""

from __future__ import annotations

from ..config import Config
from ..ir import BlockType, Document


def render(doc: Document, cfg: Config, out_path: str) -> str:
    from docx import Document as Docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Docx()
    for b in doc.ordered_blocks():
        if b.type == BlockType.TABLE and b.table and b.table.rows:
            rows = b.table.rows
            ncols = max(len(r) for r in rows)
            t = d.add_table(rows=len(rows), cols=ncols)
            t.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci in range(ncols):
                    cell_text = row[ci].output_text if ci < len(row) else ""
                    t.rows[ri].cells[ci].text = cell_text
            d.add_paragraph("")
            continue

        text = b.output_text.strip()
        if not text:
            continue

        if b.type == BlockType.TITLE:
            d.add_heading(text, level=0)
        elif b.type == BlockType.HEADING:
            d.add_heading(text, level=max(1, min(9, b.style.heading_level or 1)))
        elif b.type == BlockType.LIST_ITEM:
            d.add_paragraph(text, style="List Bullet")
        else:
            p = d.add_paragraph(text)
            if b.style.rtl:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    d.save(out_path)
    return out_path
