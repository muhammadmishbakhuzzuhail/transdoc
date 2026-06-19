# SPDX-License-Identifier: AGPL-3.0-only
# Copyright (C) 2026 Muhammad Mishbakhuz Zuhail
"""Generate synthetic test cases that real downloads don't cover cleanly.

Produces, with known ground-truth text:
  - image_only/   : PNG images of text in many scripts (the "doc is just an image" case)
  - scanned_pdf/  : image-only PDF, no text layer (the "scanned text" case)
  - photo/        : rotated + noisy image (the "phone photo" case)
  - docx/, odt/   : a structured doc with heading + table + multilingual text
Run with the project venv: .venv/bin/python scripts/make_samples.py
"""

from __future__ import annotations

from pathlib import Path

import fitz  # PyMuPDF
import numpy as np
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent.parent
S = ROOT / "corpus" / "synthetic"

NOTO = "/usr/share/fonts/noto/NotoSans-Regular.ttf"
NOTO_AR = "/usr/share/fonts/noto/NotoSansArabic-Regular.ttf"
NOTO_DEV = "/usr/share/fonts/noto/NotoSansDevanagari-Regular.ttf"
NOTO_CJK = "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc"

# Ground-truth samples: (filename, font, text, rtl)
SCRIPTS = [
    ("en", NOTO, "The quick brown fox jumps over the lazy dog. 1234567890.", False),
    ("ru", NOTO, "Съешь же ещё этих мягких французских булок да выпей чаю.", False),
    ("zh", NOTO_CJK, "快速的棕色狐狸跳过了那只懒狗。人工智能与机器翻译。", False),
    ("ja", NOTO_CJK, "いろはにほへと ちりぬるを。機械翻訳のテストです。", False),
    ("ko", NOTO_CJK, "다람쥐 헌 쳇바퀴에 타고파. 기계 번역 테스트입니다.", False),
    ("hi", NOTO_DEV, "तेज़ भूरी लोमड़ी आलसी कुत्ते के ऊपर से कूद गई।", False),
    ("ar", NOTO_AR, "العربية لغة جميلة. الترجمة الآلية اختبار النص.", True),
]


def _shape_ar(text: str) -> str:
    import arabic_reshaper
    from bidi.algorithm import get_display

    return get_display(arabic_reshaper.reshape(text))


def render_text_png(text: str, font_path: str, rtl: bool, out: Path, size: int = 40):
    if rtl:
        text = _shape_ar(text)
    font = ImageFont.truetype(font_path, size)
    dummy = Image.new("RGB", (10, 10))
    d = ImageDraw.Draw(dummy)
    bbox = d.textbbox((0, 0), text, font=font)
    w, h = bbox[2] - bbox[0] + 80, bbox[3] - bbox[1] + 80
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    x = w - bbox[2] - 40 if rtl else 40
    d.text((x, 40), text, fill="black", font=font)
    img.save(out)


def make_image_only():
    out = S / "image_only"
    out.mkdir(parents=True, exist_ok=True)
    for name, font, text, rtl in SCRIPTS:
        render_text_png(text, font, rtl, out / f"text_{name}.png")
    print(f"image_only: {len(SCRIPTS)} PNGs")


def make_scanned_pdf(src: Path, dst: Path, dpi: int = 150):
    """Rasterize a digital PDF into an image-only PDF (kills the text layer)."""
    doc = fitz.open(str(src))
    out = fitz.open()
    for page in doc:
        pix = page.get_pixmap(dpi=dpi)
        rect = fitz.Rect(0, 0, pix.width * 72 / dpi, pix.height * 72 / dpi)
        p = out.new_page(width=rect.width, height=rect.height)
        p.insert_image(rect, pixmap=pix)
    dst.parent.mkdir(parents=True, exist_ok=True)
    out.save(str(dst))
    print(f"scanned_pdf: {dst.name} ({out.page_count} pages, no text layer)")


def make_photo(src_png: Path, dst: Path, angle: float = -7.0):
    """Simulate a phone photo: rotate, add noise + uneven lighting."""
    img = Image.open(src_png).convert("RGB")
    img = img.rotate(angle, expand=True, fillcolor=(245, 243, 238))
    arr = np.asarray(img).astype(np.float32)
    # uneven lighting gradient
    h, w = arr.shape[:2]
    grad = np.linspace(0.7, 1.05, w)[None, :, None]
    arr = np.clip(arr * grad, 0, 255)
    # gaussian noise
    arr = np.clip(arr + np.random.default_rng(0).normal(0, 8, arr.shape), 0, 255)
    dst.parent.mkdir(parents=True, exist_ok=True)
    Image.fromarray(arr.astype(np.uint8)).save(dst, quality=70)
    print(f"photo: {dst.name} (rotated {angle}deg + noise + lighting)")


def make_docx(dst: Path):
    from docx import Document as Docx

    d = Docx()
    d.add_heading("Document Intelligence Test", level=1)
    d.add_paragraph("This DOCX exercises headings, paragraphs, lists, and a table.")
    d.add_heading("Multilingual paragraph", level=2)
    d.add_paragraph("English. Русский текст. 中文文本。 العربية. हिन्दी पाठ।")
    d.add_paragraph("First bullet", style="List Bullet")
    d.add_paragraph("Second bullet", style="List Bullet")
    t = d.add_table(rows=3, cols=3)
    hdr = t.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "ID", "Name", "Amount"
    data = [("001", "Invoice A", "1.250,00"), ("002", "Invoice B", "980,50")]
    for r, (a, b, c) in enumerate(data, start=1):
        cells = t.rows[r].cells
        cells[0].text, cells[1].text, cells[2].text = a, b, c
    dst.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(dst))
    print(f"docx: {dst.name}")


if __name__ == "__main__":
    make_image_only()
    make_scanned_pdf(S / "multilingual" / "udhr_english.pdf", S / "scanned_pdf" / "udhr_english_scanned.pdf")
    make_scanned_pdf(S / "multilingual" / "udhr_russian.pdf", S / "scanned_pdf" / "udhr_russian_scanned.pdf")
    make_photo(S / "image_only" / "text_en.png", S / "photo" / "photo_en.jpg")
    make_docx(S / "docx" / "structured.docx")
    print("\nDONE. Convert docx->odt with: soffice --headless --convert-to odt --outdir corpus/synthetic/odt corpus/synthetic/docx/structured.docx")
