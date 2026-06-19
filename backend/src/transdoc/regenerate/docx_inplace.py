# © 2026 Muhammad Mishbakhuz Zuhail. All rights reserved.
# Proprietary — source-available for reference only; no use, copying, or
# distribution without written permission. See LICENSE.
"""In-place DOCX translation — the DeepL strategy for Office files.

Re-open the *source* .docx and swap the translated text into the existing paragraphs and
table cells, leaving the document structure untouched. Because we only change run text, every
style, image, header/footer, table, list and section stays exactly as authored, and Word's
own layout engine handles the (longer) translated text on open — no overlay, no rebuild.

When the block carries inline runs (a paragraph whose text is NOT uniformly styled — a bold word,
a superscript ref, an inline link), the per-run translations are written back as styled runs so
word-level formatting survives. A uniformly-styled paragraph keeps its run-0 formatting for the
whole translation (no rebuild). Paragraph-level style and everything else stay untouched.
"""

from __future__ import annotations

from ..config import Config
from ..extract.docx import iter_block_items
from ..ir import Block, Document


def _set_paragraph(paragraph, b: Block) -> None:
    from .docx_out import _set_para_bidi, _set_run_rtl
    # An LTR source paragraph translated into an RTL target must flip its base direction.
    if b.style.rtl:
        _set_para_bidi(paragraph)
    if getattr(b, "runs", None):
        # mixed-style paragraph: blank the source runs, re-emit the translated runs WITH their
        # captured character style (bold/italic/super/link/...) — preserves inline formatting a
        # flatten-to-run-0 would lose.
        from .docx_out import _add_run
        for r in list(paragraph.runs):
            r.text = ""
        for run in b.runs:
            _add_run(paragraph, run)   # _add_run sets run-level rtl from run.style.rtl
        return
    runs = paragraph.runs
    if not runs:
        r = paragraph.add_run(b.output_text)
        if b.style.rtl:
            _set_run_rtl(r)
        return
    runs[0].text = b.output_text        # uniform paragraph: keep run-0 formatting
    if b.style.rtl:
        _set_run_rtl(runs[0])
    for r in runs[1:]:
        r.text = ""


def _set_cell(cell, text: str) -> None:
    paras = cell.paragraphs
    if paras:
        runs = paras[0].runs
        if runs:
            runs[0].text = text
            for r in runs[1:]:
                r.text = ""
        else:
            paras[0].add_run(text)
        for p in paras[1:]:
            for r in p.runs:
                r.text = ""
    else:
        cell.text = text


def render(doc: Document, cfg: Config, out_path: str) -> str:
    from docx import Document as Docx
    from docx.text.paragraph import Paragraph

    d = Docx(doc.source_path)
    blocks = doc.ordered_blocks()

    # In-place editing pairs each source block-item (non-empty paragraph / table) with the IR
    # block at the same index. That holds ONLY when extraction emitted exactly one block per item
    # in order — but furniture/header stripping, fuse reconcile, or caption reordering can drop or
    # merge blocks, after which index-zipping writes every later translation into the WRONG
    # paragraph (silent, catastrophic — audit finding). When the counts diverge, fall back to the
    # IR-rebuild renderer: it loses some in-place formatting but stays correct.
    items = [it for it in iter_block_items(d)
             if not (isinstance(it, Paragraph) and not it.text.strip())]
    if len(items) != len(blocks):
        from .docx_out import render as render_out
        return render_out(doc, cfg, out_path)

    for item, b in zip(items, blocks):
        if isinstance(item, Paragraph):
            _set_paragraph(item, b)
        elif b.table:  # table — walk cells in the same row/col order the extractor used
            tcells = [c for row in b.table.rows for c in row]
            dcells = [c for row in item.rows for c in row.cells]
            for dc, tc in zip(dcells, tcells):
                if tc.text.strip():
                    _set_cell(dc, tc.output_text)

    d.save(out_path)
    return out_path
