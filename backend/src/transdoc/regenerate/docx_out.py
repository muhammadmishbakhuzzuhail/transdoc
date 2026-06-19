# © 2026 Muhammad Mishbakhuz Zuhail. All rights reserved.
# Proprietary — source-available for reference only; no use, copying, or
# distribution without written permission. See LICENSE.
"""DOCX renderer (flow fidelity). Rebuilds an editable Word doc from the IR.

Layout = logical structure (headings, paragraphs, lists, tables), not pixel-exact — the
honest trade-off for an *editable* target. For visual fidelity use the PDF overlay renderer.
"""

from __future__ import annotations

from ..config import Config
from ..ir import BlockType, Document, Style
from ..textdir import is_rtl_text


def _align(style: Style):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if style.rtl:
        return WD_ALIGN_PARAGRAPH.RIGHT
    return {"center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "left": WD_ALIGN_PARAGRAPH.LEFT}.get(style.align)


def _style_runs(para, style: Style) -> None:
    """Apply the IR character style (font/size/weight/colour) to a paragraph's runs — the IR
    captures these but the renderer used to drop them (plain 12pt black). Heading sizes are left
    to the Word style unless the source carried an explicit size."""
    from docx.shared import Pt, RGBColor

    rgb = None
    if style.color and style.color.startswith("#") and len(style.color) == 7:
        try:
            rgb = RGBColor(int(style.color[1:3], 16), int(style.color[3:5], 16),
                           int(style.color[5:7], 16))
        except ValueError:
            rgb = None
    for run in para.runs:
        f = run.font
        if style.bold:
            f.bold = True
        if style.italic:
            f.italic = True
        if style.underline:
            f.underline = True
        if style.strike:
            f.strike = True
        if style.small_caps:
            f.small_caps = True
        if style.font:
            f.name = style.font
        if style.size and style.size > 0:
            f.size = Pt(style.size)
        if rgb is not None:
            f.color.rgb = rgb
        if style.rtl:
            _set_run_rtl(run)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Append an external hyperlink run (python-docx has no public API). Blue + underlined so
    it reads as a link."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def _set_para_bidi(p) -> None:
    """Mark a paragraph right-to-left (pPr/w:bidi) so Word runs the bidi algorithm + flows it RTL.
    Alignment is set separately via _align; this sets the *base direction*."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    ppr = p._p.get_or_add_pPr()
    if ppr.find(qn("w:bidi")) is None:
        ppr.append(OxmlElement("w:bidi"))


def _set_run_rtl(run) -> None:
    """Mark a run right-to-left (rPr/w:rtl) so its characters + adjacent punctuation order RTL."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    rpr = run._r.get_or_add_rPr()
    if rpr.find(qn("w:rtl")) is None:
        rpr.append(OxmlElement("w:rtl"))


def _apply_para_format(p, style) -> None:
    """Carry paragraph spacing + indentation onto the output paragraph."""
    from docx.shared import Pt
    if style.rtl:
        _set_para_bidi(p)
    pf = p.paragraph_format
    if style.space_before is not None:
        pf.space_before = Pt(style.space_before)
    if style.space_after is not None:
        pf.space_after = Pt(style.space_after)
    if style.line_spacing:
        pf.line_spacing = style.line_spacing
    if style.indent_left is not None:
        pf.left_indent = Pt(style.indent_left)
    if style.indent_first is not None:
        pf.first_line_indent = Pt(style.indent_first)
    if style.para_shading or style.para_border:
        _apply_para_shading_border(p, style)
    if style.tab_stops:
        try:
            from docx.enum.text import WD_TAB_ALIGNMENT
            amap = {"left": WD_TAB_ALIGNMENT.LEFT, "center": WD_TAB_ALIGNMENT.CENTER,
                    "right": WD_TAB_ALIGNMENT.RIGHT, "decimal": WD_TAB_ALIGNMENT.DECIMAL}
            for pos, al in style.tab_stops:
                pf.tab_stops.add_tab_stop(Pt(pos), amap.get(al, WD_TAB_ALIGNMENT.LEFT))
        except Exception:
            pass
    if style.drop_cap:
        _apply_drop_cap(p)


def _apply_drop_cap(p) -> None:
    """Reproduce a drop-cap by adding a pPr/framePr dropCap frame (best effort)."""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        ppr = p._p.get_or_add_pPr()
        fp = OxmlElement("w:framePr")
        fp.set(qn("w:dropCap"), "drop")
        fp.set(qn("w:lines"), "3")
        fp.set(qn("w:wrap"), "around")
        fp.set(qn("w:vAnchor"), "text")
        fp.set(qn("w:hAnchor"), "text")
        ppr.append(fp)
    except Exception:
        pass


def _apply_para_shading_border(p, style) -> None:
    """Reproduce a boxed/callout paragraph: pPr/w:shd background fill + pPr/w:pBdr box border."""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        ppr = p._p.get_or_add_pPr()
        if style.para_shading:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), style.para_shading.lstrip("#"))
            ppr.append(shd)
        if style.para_border:
            pbdr = OxmlElement("w:pBdr")
            for edge in ("top", "left", "bottom", "right"):
                e = OxmlElement(f"w:{edge}")
                e.set(qn("w:val"), "single")
                e.set(qn("w:sz"), "4")
                e.set(qn("w:space"), "4")
                e.set(qn("w:color"), "auto")
                pbdr.append(e)
            ppr.append(pbdr)
    except Exception:
        pass


def _add_run(p, run) -> None:
    """Add one inline run with its own style (bold/italic/underline/super/sub/size/colour),
    or a hyperlink run when the run is a link."""
    from docx.shared import Pt, RGBColor

    s = run.style
    if s.link:
        _add_hyperlink(p, s.link, run.output_text)
        return
    r = p.add_run(run.output_text)
    f = r.font
    f.bold = s.bold or None
    f.italic = s.italic or None
    f.underline = s.underline or None
    f.strike = s.strike or None
    if s.small_caps:
        f.small_caps = True
    if s.all_caps:
        f.all_caps = True
    if s.highlight:
        try:
            from docx.enum.text import WD_COLOR_INDEX
            f.highlight_color = WD_COLOR_INDEX[s.highlight.upper()]
        except Exception:
            pass
    if s.superscript:
        f.superscript = True
    if s.subscript:
        f.subscript = True
    if s.font:
        f.name = s.font
    if s.size and s.size > 0:
        f.size = Pt(s.size)
    if s.color and s.color.startswith("#") and len(s.color) == 7:
        try:
            f.color.rgb = RGBColor(int(s.color[1:3], 16), int(s.color[3:5], 16),
                                   int(s.color[5:7], 16))
        except ValueError:
            pass
    if s.rtl:
        _set_run_rtl(r)


def _render_table(d, table) -> None:
    """Build a DOCX table honoring merged cells + column widths. The IR stores a spanning cell
    once with colspan/rowspan (HTML semantics), so rows can have fewer cells than the grid is
    wide; place each cell at the next free grid slot and merge across its span."""
    from docx.shared import Pt
    rows = table.rows
    ncols = max((sum(max(1, c.colspan) for c in r) for r in rows), default=1)
    nrows = len(rows)
    t = d.add_table(nrows, ncols)        # positional: works for Document + _Cell (nested)
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    if table.col_widths:
        t.autofit = False
        for i, w in enumerate(table.col_widths):
            if w and w > 0 and i < ncols:
                for cell in t.columns[i].cells:
                    cell.width = Pt(w)
    if table.row_heights:
        for i, h in enumerate(table.row_heights):
            if h and h > 0 and i < len(t.rows):
                t.rows[i].height = Pt(h)
    if table.cell_margin and table.cell_margin > 0:
        _set_table_cell_margin(t, table.cell_margin)
    occupied: set[tuple[int, int]] = set()
    for ri, row in enumerate(rows):
        ci = 0
        for cell in row:
            while (ri, ci) in occupied and ci < ncols - 1:
                ci += 1
            r2 = min(ri + max(1, cell.rowspan) - 1, nrows - 1)
            c2 = min(ci + max(1, cell.colspan) - 1, ncols - 1)
            anchor = t.cell(ri, ci)
            if r2 > ri or c2 > ci:
                try:
                    anchor = anchor.merge(t.cell(r2, c2))
                except Exception:
                    pass
            anchor.text = cell.output_text
            if cell.shading:
                try:
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:fill"), cell.shading.lstrip("#"))
                    anchor._tc.get_or_add_tcPr().append(shd)
                except Exception:
                    pass
            if (cell.size or cell.bold) and anchor.paragraphs and anchor.paragraphs[0].runs:
                from docx.shared import Pt
                run = anchor.paragraphs[0].runs[0]
                if cell.size and cell.size > 0:
                    run.font.size = Pt(cell.size)
                if cell.bold:
                    run.font.bold = True
            if cell.align and anchor.paragraphs:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                anchor.paragraphs[0].alignment = {
                    "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "left": WD_ALIGN_PARAGRAPH.LEFT}.get(cell.align)
            # RTL cell text -> base direction so Word flows + orders it right-to-left
            if anchor.paragraphs and is_rtl_text(cell.output_text):
                _set_para_bidi(anchor.paragraphs[0])
                for rr in anchor.paragraphs[0].runs:
                    _set_run_rtl(rr)
            if cell.table is not None:
                _render_table(anchor, cell.table)   # nested table inside the cell
            for rr in range(ri, r2 + 1):
                for cc in range(ci, c2 + 1):
                    occupied.add((rr, cc))
            ci = c2 + 1


def _set_table_cell_margin(t, pt) -> None:
    """Set a uniform tblCellMar (cell padding) on all four edges of the table."""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        twips = str(int(pt * 20))
        tblpr = t._tbl.tblPr
        mar = OxmlElement("w:tblCellMar")
        for edge in ("top", "left", "bottom", "right"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:w"), twips)
            e.set(qn("w:type"), "dxa")
            mar.append(e)
        tblpr.append(mar)
    except Exception:
        pass


def _set_doc_language(d, lang) -> None:
    """Tag the rebuilt document's default language as the TARGET language so Word's spell-check
    and hyphenation match the translation (the source per-run lang tags no longer apply — every
    run is now target-language text). Sets w:lang on the Normal style's run properties."""
    if not lang:
        return
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        code = str(lang).replace("_", "-")
        rpr = d.styles["Normal"].element.get_or_add_rPr()
        for tag in rpr.findall(qn("w:lang")):
            rpr.remove(tag)
        el = OxmlElement("w:lang")
        el.set(qn("w:val"), code)
        el.set(qn("w:bidi"), code)
        rpr.append(el)
    except Exception:
        pass


def _list_style(d, b):
    """Word list style honoring nesting: 'List Bullet'/'List Number' for level 0, then the
    leveled variants ('List Bullet 2'..'3') for deeper items. Falls back to the base style (then
    None) when a template lacks the leveled one — python-docx raises KeyError on a missing style."""
    from ..ir import BlockType
    if b.type != BlockType.LIST_ITEM:
        return None
    base = "List Number" if b.style.list_ordered else "List Bullet"
    lvl = max(0, min(b.style.list_level or 0, 2))     # built-in templates cover base + 2/3
    names = ([f"{base} {lvl + 1}"] if lvl else []) + [base]
    avail = {s.name for s in d.styles}
    for n in names:
        if n in avail:
            return n
    return None


def _fill_hf(part, blocks) -> None:
    """Write captured header/footer blocks (translated) into a section's header or footer part.
    The default part has one empty paragraph; reuse it for the first block, add the rest."""
    if not blocks:
        return
    part.is_linked_to_previous = False
    existing = part.paragraphs
    for i, b in enumerate(blocks):
        p = existing[0] if (i == 0 and existing) else part.add_paragraph()
        p.text = b.output_text.strip()
        _style_runs(p, b.style)
        if b.style.rtl:
            _set_para_bidi(p)
        align = _align(b.style)
        if align is not None:
            p.alignment = align


def render(doc: Document, cfg: Config, out_path: str) -> str:
    from docx import Document as Docx
    from docx.shared import Inches

    d = Docx()
    _set_doc_language(d, doc.target_lang or cfg.target_lang)
    for b in doc.ordered_blocks():
        if b.page_break_before:
            d.add_page_break()           # reproduce the source's manual page break
        if b.type == BlockType.FIGURE and b.image_path:
            try:
                d.add_picture(b.image_path, width=Inches(5.5))
            except Exception:
                pass
            continue

        # Formula: a verbatim crop (pixel-perfect math) when available, else the LaTeX text.
        if b.type == BlockType.FORMULA and b.image_path:
            try:
                w = (b.bbox.x1 - b.bbox.x0) / 72.0 if b.bbox else 3.0
                d.add_picture(b.image_path, width=Inches(max(0.5, min(5.5, w))))
            except Exception:
                d.add_paragraph(b.output_text.strip())
            continue

        if b.type == BlockType.TABLE and b.table and b.table.rows:
            _render_table(d, b.table)
            d.add_paragraph("")
            continue

        text = b.output_text.strip()
        if not text:
            continue

        # bilingual: source (italic) then translation, mirroring the markdown renderer
        if cfg.bilingual and b.translated is not None and b.text.strip():
            src_p = d.add_paragraph()
            src_p.add_run(b.text.strip()).italic = True
            d.add_paragraph(b.translated.strip())
            continue

        if b.runs and b.type in (BlockType.PARAGRAPH, BlockType.CAPTION, BlockType.LIST_ITEM):
            p = d.add_paragraph(style=_list_style(d, b))
            for run in b.runs:
                _add_run(p, run)         # inline styled runs (bold word / superscript / link)
            align = _align(b.style)
            if align is not None:
                p.alignment = align
            _apply_para_format(p, b.style)
            continue
        if b.style.link and b.type in (BlockType.PARAGRAPH, BlockType.CAPTION,
                                       BlockType.LIST_ITEM):
            p = d.add_paragraph(style=_list_style(d, b))
            _add_hyperlink(p, b.style.link, text)   # link run, not a plain run
        elif b.type == BlockType.TITLE:
            p = d.add_heading(text, level=0)
        elif b.type == BlockType.HEADING:
            p = d.add_heading(text, level=max(1, min(9, b.style.heading_level or 1)))
        elif b.type == BlockType.LIST_ITEM:
            p = d.add_paragraph(text, style=_list_style(d, b))
        else:
            p = d.add_paragraph(text)
        # carry the source character styling + alignment into the output (was dropped)
        _style_runs(p, b.style)
        align = _align(b.style)
        if align is not None:
            p.alignment = align
        _apply_para_format(p, b.style)

    if doc.headers or doc.footers:
        try:
            sec0 = d.sections[0]
            _fill_hf(sec0.header, doc.headers)
            _fill_hf(sec0.footer, doc.footers)
        except Exception:
            pass

    if doc.section_columns and doc.section_columns > 1:
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            sectpr = d.sections[0]._sectPr
            for old in sectpr.findall(qn("w:cols")):
                sectpr.remove(old)
            cols = OxmlElement("w:cols")
            cols.set(qn("w:num"), str(doc.section_columns))
            cols.set(qn("w:space"), "425")        # 0.3" gutter (twips)
            sectpr.append(cols)
        except Exception:
            pass

    pm = doc.page_margins or {}
    if pm:
        from docx.shared import Pt
        try:
            sec = d.sections[0]
            if pm.get("left"):
                sec.left_margin = Pt(pm["left"])
            if pm.get("right"):
                sec.right_margin = Pt(pm["right"])
            if pm.get("top"):
                sec.top_margin = Pt(pm["top"])
            if pm.get("bottom"):
                sec.bottom_margin = Pt(pm["bottom"])
        except Exception:
            pass

    md = doc.metadata or {}
    cp = d.core_properties
    if md.get("title"):
        cp.title = md["title"]
    if md.get("author"):
        cp.author = md["author"]
    if md.get("subject"):
        cp.subject = md["subject"]
    if md.get("keywords"):
        cp.keywords = md["keywords"]
    d.save(out_path)
    return out_path
