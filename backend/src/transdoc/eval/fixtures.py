"""Deterministic, digital-only eval fixtures for the CI regression gate.

CI can't run the OCR/paddle paths (non-deterministic, heavy), so the committed baseline is
built from these born-digital documents whose extraction + echo render are byte-stable across
machines. Regenerate the corpus + baseline with:

    python -m transdoc.eval.fixtures src/transdoc/eval/samples
    python -m transdoc.eval.harness src/transdoc/eval/samples --engine echo \
        --out src/transdoc/eval/baseline.json
"""

from __future__ import annotations

import sys
from pathlib import Path


def _pdf(path: Path) -> None:
    import fitz

    d = fitz.open()
    for i in range(2):
        p = d.new_page(width=595, height=842)
        p.insert_text((42, 60), f"Section {i + 1}: a heading line", fontsize=16)
        p.insert_text((42, 100), "A body paragraph with enough words to translate cleanly.",
                      fontsize=11)
        p.insert_text((42, 130), "A second body paragraph on the same page, also translatable.",
                      fontsize=11)
    d.save(str(path))
    d.close()


def _docx(path: Path) -> None:
    from docx import Document as Docx

    d = Docx()
    d.add_heading("A document title", level=0)
    d.add_paragraph("An introductory paragraph that should be translated by the echo engine.")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Name"
    t.cell(0, 1).text = "Value"
    t.cell(1, 0).text = "alpha"
    t.cell(1, 1).text = "beta"
    d.add_paragraph("A closing paragraph after the table.")
    d.save(str(path))


def build(out_dir: str | Path) -> list[Path]:
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    pdf, docx = out / "digital_two_page.pdf", out / "digital_table.docx"
    _pdf(pdf)
    _docx(docx)
    return [pdf, docx]


def main() -> None:
    out = sys.argv[1] if len(sys.argv) > 1 else "src/transdoc/eval/samples"
    built = build(out)
    for p in built:
        print(f"wrote {p}")


if __name__ == "__main__":
    main()
