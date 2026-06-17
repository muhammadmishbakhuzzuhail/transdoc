"""In-place DOCX translation (DeepL strategy): a .docx source is mutated, swapping text while
keeping every paragraph style, list, and table — not rebuilt from the IR."""

from __future__ import annotations

import pytest

pytest.importorskip("docx")

from docx import Document as Docx  # noqa: E402

from transdoc.config import Config, Engine, Mode, OutputFormat  # noqa: E402
from transdoc.pipeline import run  # noqa: E402


def _make_docx(path):
    d = Docx()
    d.add_heading("Project Report", level=1)
    d.add_paragraph("This is the body paragraph with some content.")
    d.add_paragraph("First item", style="List Bullet")
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Name"
    t.rows[0].cells[1].text = "Value"
    t.rows[1].cells[0].text = "Alpha"
    t.rows[1].cells[1].text = "One"
    d.save(str(path))


def test_inplace_keeps_styles_and_swaps_text(tmp_path):
    src = tmp_path / "doc.docx"
    _make_docx(src)
    out = tmp_path / "doc.id.docx"
    run(str(src), Config(source_lang="en", target_lang="id", engine=Engine.ECHO,
                         output_format=OutputFormat.DOCX, mode=Mode.FULL), out_path=str(out))

    o, t = Docx(str(src)), Docx(str(out))
    # structure preserved 1:1
    assert len(o.paragraphs) == len(t.paragraphs)
    assert len(o.tables) == len(t.tables)
    # styles preserved, text translated (echo prefixes "[id] ")
    by_style = {p.style.name: p.text for p in t.paragraphs if p.text.strip()}
    assert any("Heading" in s and v.startswith("[id]") for s, v in by_style.items())
    assert any("List" in s for s in by_style)
    # table cell translated in place
    assert t.tables[0].rows[0].cells[0].text == "[id] Name"


def test_inplace_preserves_embedded_image(tmp_path):
    """In-place editing only rewrites run text, so an embedded image must survive untouched."""
    from PIL import Image
    img = tmp_path / "logo.png"
    Image.new("RGB", (32, 32), (200, 30, 30)).save(str(img))
    src = tmp_path / "doc.docx"
    d = Docx()
    d.add_heading("Report", level=1)
    d.add_paragraph("Body text to translate here.")
    d.add_picture(str(img))
    d.save(str(src))

    out = tmp_path / "doc.id.docx"
    run(str(src), Config(source_lang="en", target_lang="id", engine=Engine.ECHO,
                         output_format=OutputFormat.DOCX, mode=Mode.FULL), out_path=str(out))
    o, t = Docx(str(src)), Docx(str(out))
    assert len(t.inline_shapes) == len(o.inline_shapes) >= 1   # image preserved in place


def test_inplace_falls_back_when_block_count_diverges(tmp_path, monkeypatch):
    """If extraction/reconcile dropped a block, the index-zip would misalign every later
    paragraph. The renderer must detect the count mismatch and fall back to the IR rebuild
    instead of writing translations into the wrong paragraphs."""
    from transdoc.config import Config
    from transdoc.extract import extract
    from transdoc.ingest.detect import detect
    from transdoc.regenerate import docx_inplace, docx_out

    src = tmp_path / "doc.docx"
    _make_docx(src)
    cfg = Config(source_lang="en", target_lang="id", engine=Engine.ECHO)
    doc = extract(detect(str(src)), cfg)
    for b in doc.blocks:          # echo-translate so output_text is set
        if b.is_translatable:
            b.translated = b.text
    doc.blocks = doc.blocks[:-1]  # simulate a dropped block -> count now diverges from source

    called = {"out": False}
    real_out = docx_out.render
    monkeypatch.setattr(docx_out, "render",
                        lambda *a, **k: called.__setitem__("out", True) or real_out(*a, **k))
    out = tmp_path / "doc.id.docx"
    docx_inplace.render(doc, cfg, str(out))
    assert called["out"] is True   # delegated to the safe rebuild path
    assert out.exists()
