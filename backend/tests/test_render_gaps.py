# SPDX-License-Identifier: AGPL-3.0-only
# Copyright (C) 2026 Muhammad Mishbakhuz Zuhail
"""Render-gap fixes: table cell alignment + list nesting (docx), PDF output /Lang."""

from __future__ import annotations

import pytest

from transdoc.config import Config
from transdoc.ir import BBox, Block, BlockType, Cell, Confidence, Document, Style, Table

docx = pytest.importorskip("docx")


def test_docx_cell_alignment_applied(tmp_path):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    from transdoc.regenerate.docx_out import render
    tbl = Table(rows=[[Cell(text="L", align="left"), Cell(text="R", align="right"),
                       Cell(text="C", align="center")]])
    d = Document(source_path="x", mime="docx")
    d.blocks = [Block(id="t", type=BlockType.TABLE, bbox=BBox(x0=0, y0=0, x1=1, y1=1),
                      table=tbl, confidence=Confidence())]
    out = tmp_path / "o.docx"
    render(d, Config(target_lang="id"), str(out))
    row = docx.Document(str(out)).tables[0].rows[0]
    assert row.cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert row.cells[2].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_docx_list_nesting_uses_leveled_style(tmp_path):
    from transdoc.regenerate.docx_out import render
    d = Document(source_path="x", mime="docx")
    d.blocks = [
        Block(id="a", type=BlockType.LIST_ITEM, text="top",
              style=Style(list_level=0), confidence=Confidence()),
        Block(id="b", type=BlockType.LIST_ITEM, text="nested",
              style=Style(list_level=1), confidence=Confidence()),
    ]
    out = tmp_path / "o.docx"
    render(d, Config(target_lang="id"), str(out))
    styles = [p.style.name for p in docx.Document(str(out)).paragraphs if p.text.strip()]
    assert "List Bullet" in styles[0]
    assert styles[1].endswith("2")          # nested -> 'List Bullet 2'


def test_pdf_output_lang_tag(tmp_path):
    pytest.importorskip("fitz")
    from transdoc.regenerate.pdf_out import render_reconstruct
    d = Document(source_path="x", mime="application/pdf", target_lang="id")
    d.page_sizes[0] = (300, 200)
    d.blocks = [Block(id="a", type=BlockType.PARAGRAPH, text="halo",
                      bbox=BBox(x0=10, y0=10, x1=200, y1=30), translated="halo",
                      confidence=Confidence())]
    out = tmp_path / "o.pdf"
    render_reconstruct(d, Config(target_lang="id"), str(out))
    raw = open(out, "rb").read()
    assert b"/Lang" in raw and b"id" in raw
