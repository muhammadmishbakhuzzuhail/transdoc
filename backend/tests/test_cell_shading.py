"""Table cell background shading captured (docx) + rendered (docx + pdf)."""

from __future__ import annotations

import pytest

from transdoc.ir import Cell


def test_pdf_cell_shading():
    from transdoc.regenerate.pdf_out import _cell_td
    assert "background-color:#ffcc00" in _cell_td(Cell(text="x", shading="#ffcc00"))


def test_docx_extract_and_render_cell_shading(tmp_path):
    docx = pytest.importorskip("docx")
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from transdoc.extract.docx import extract
    from transdoc.ir import BlockType
    from transdoc.regenerate.docx_out import render
    from transdoc.config import Config as C

    dd = docx.Document()
    t = dd.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    cell.text = "shaded"
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "FFFF00")
    cell._tc.get_or_add_tcPr().append(shd)
    src = tmp_path / "s.docx"
    dd.save(str(src))

    doc = extract(str(src), C(target_lang="id"))
    blk = next(b for b in doc.blocks if b.type == BlockType.TABLE)
    assert blk.table.rows[0][0].shading == "#FFFF00"

    out = tmp_path / "o.docx"
    render(doc, C(target_lang="id"), str(out))
    rt = docx.Document(str(out)).tables[0]
    tcpr = rt.cell(0, 0)._tc.tcPr
    assert tcpr is not None and tcpr.find(qn("w:shd")) is not None
