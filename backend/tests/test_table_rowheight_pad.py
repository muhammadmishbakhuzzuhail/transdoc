"""Table row heights + uniform cell padding (margin): captured from DOCX, rendered docx + pdf."""

from __future__ import annotations

import pytest

from transdoc.config import Config

docx = pytest.importorskip("docx")


def test_capture_row_height_and_margin(tmp_path):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    from transdoc.extract.docx import extract
    from transdoc.ir import BlockType
    dd = docx.Document()
    t = dd.add_table(rows=2, cols=2)
    t.rows[0].height = Pt(40)
    t.cell(0, 0).text = "a"
    # set a tblCellMar top = 100 twips (5pt)
    tblpr = t._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    top = OxmlElement("w:top")
    top.set(qn("w:w"), "100")
    top.set(qn("w:type"), "dxa")
    mar.append(top)
    tblpr.append(mar)
    f = tmp_path / "in.docx"
    dd.save(str(f))
    blk = next(b for b in extract(str(f), Config(target_lang="id")).blocks
               if b.type == BlockType.TABLE)
    assert blk.table.row_heights and round(blk.table.row_heights[0]) == 40
    assert blk.table.cell_margin and round(blk.table.cell_margin) == 5


def _tbl_doc():
    from transdoc.ir import BBox, Block, BlockType, Cell, Confidence, Document, Table
    tbl = Table(rows=[[Cell(text="x"), Cell(text="y")]],
                row_heights=[40.0], cell_margin=5.0, col_widths=[100.0, 100.0])
    d = Document(source_path="x", mime="docx")
    d.blocks = [Block(id="t", type=BlockType.TABLE, bbox=BBox(x0=0, y0=0, x1=1, y1=1),
                      table=tbl, confidence=Confidence())]
    return d


def test_docx_render_row_height_margin(tmp_path):
    from docx.oxml.ns import qn

    from transdoc.regenerate.docx_out import render
    out = tmp_path / "o.docx"
    render(_tbl_doc(), Config(target_lang="id"), str(out))
    t = docx.Document(str(out)).tables[0]
    assert t.rows[0].height is not None
    assert t._tbl.tblPr.find(qn("w:tblCellMar")) is not None


def test_pdf_html_row_height_padding():
    from transdoc.regenerate.pdf_out import _table_html
    html = _table_html(_tbl_doc().blocks[0].table)
    assert "height:40pt" in html and "padding:5pt" in html
