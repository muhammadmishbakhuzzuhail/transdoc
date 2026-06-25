# SPDX-License-Identifier: AGPL-3.0-only
# Copyright (C) 2026 Muhammad Mishbakhuz Zuhail
"""Highlight + small-caps captured (docx) + rendered (md/docx/pdf), run-level."""

from __future__ import annotations

import pytest

from transdoc.config import Config
from transdoc.ir import BBox, Block, BlockType, Confidence, Document, Run, Style


def _blk(style):
    d = Document(source_path="x", mime="docx")
    d.blocks = [Block(id="1", type=BlockType.PARAGRAPH, text="word", bbox=BBox(x0=0,y0=0,x1=1,y1=1),
                      confidence=Confidence(), runs=[Run(text="word", style=style)])]
    return d


def test_markdown_highlight_smallcaps():
    from transdoc.regenerate.markdown import render
    md = render(_blk(Style(highlight="yellow", small_caps=True)), Config(target_lang="id"))
    assert "<mark>" in md and "small-caps" in md


def test_pdf_run_span_highlight_smallcaps():
    from transdoc.regenerate.pdf_out import _run_span
    h = _run_span(Run(text="x", style=Style(highlight="green", small_caps=True)))
    assert "background-color:#008000" in h and "font-variant:small-caps" in h


def test_docx_highlight_smallcaps(tmp_path):
    docx = pytest.importorskip("docx")
    from transdoc.regenerate.docx_out import render
    out = tmp_path / "h.docx"
    render(_blk(Style(highlight="yellow", small_caps=True)), Config(target_lang="id"), str(out))
    p = next(p for p in docx.Document(str(out)).paragraphs if p.text.strip())
    r = p.runs[0]
    assert r.font.small_caps and r.font.highlight_color is not None


def test_docx_extract_captures_highlight(tmp_path):
    docx = pytest.importorskip("docx")
    from docx.enum.text import WD_COLOR_INDEX

    from transdoc.extract.docx import extract
    dd = docx.Document()
    para = dd.add_paragraph("plain ")
    r = para.add_run("hi")
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r.font.small_caps = True
    p = tmp_path / "x.docx"
    dd.save(str(p))
    blk = next(b for b in extract(str(p), Config(target_lang="id")).blocks
               if b.type == BlockType.PARAGRAPH)
    assert any(rr.style.highlight == "yellow" and rr.style.small_caps for rr in blk.runs)
