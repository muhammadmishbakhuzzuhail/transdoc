"""Output DOCX is tagged with the target language (spell-check / hyphenation match the
translation; source per-run lang tags no longer apply once every run is target text)."""

from __future__ import annotations

import pytest

from transdoc.config import Config

docx = pytest.importorskip("docx")


def test_normal_style_language_is_target(tmp_path):
    from docx.oxml.ns import qn

    from transdoc.ir import Block, BlockType, Confidence, Document
    from transdoc.regenerate.docx_out import render
    d = Document(source_path="x", mime="docx", target_lang="id")
    d.blocks = [Block(id="a", type=BlockType.PARAGRAPH, text="halo dunia",
                      confidence=Confidence())]
    out = tmp_path / "o.docx"
    render(d, Config(target_lang="id"), str(out))
    rpr = docx.Document(str(out)).styles["Normal"].element.find(qn("w:rPr"))
    lang = rpr.find(qn("w:lang")) if rpr is not None else None
    assert lang is not None and lang.get(qn("w:val")) == "id"


def test_no_target_no_crash(tmp_path):
    from transdoc.ir import Block, BlockType, Confidence, Document
    from transdoc.regenerate.docx_out import render
    d = Document(source_path="x", mime="docx")
    d.blocks = [Block(id="a", type=BlockType.PARAGRAPH, text="x", confidence=Confidence())]
    out = tmp_path / "o.docx"
    render(d, Config(target_lang="id"), str(out))   # cfg target used as fallback
    assert out.exists()
