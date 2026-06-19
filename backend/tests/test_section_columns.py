# © 2026 Muhammad Mishbakhuz Zuhail. All rights reserved.
# Proprietary — source-available for reference only; no use, copying, or
# distribution without written permission. See LICENSE.
"""Multi-column section (DOCX w:cols): column count captured and re-emitted."""

from __future__ import annotations

import pytest

from transdoc.config import Config

docx = pytest.importorskip("docx")


def _num_cols(document):
    from docx.oxml.ns import qn
    cols = document.sections[0]._sectPr.find(qn("w:cols"))
    return int(cols.get(qn("w:num"))) if cols is not None and cols.get(qn("w:num")) else 1


def test_capture_two_columns(tmp_path):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from transdoc.extract.docx import extract
    dd = docx.Document()
    dd.add_paragraph("two column body")
    sectpr = dd.sections[0]._sectPr
    for old in sectpr.findall(qn("w:cols")):
        sectpr.remove(old)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2")
    sectpr.append(cols)
    f = tmp_path / "in.docx"
    dd.save(str(f))
    doc = extract(str(f), Config(target_lang="id"))
    assert doc.section_columns == 2


def test_single_column_default(tmp_path):
    from transdoc.extract.docx import extract
    dd = docx.Document()
    dd.add_paragraph("plain")
    f = tmp_path / "s.docx"
    dd.save(str(f))
    assert extract(str(f), Config(target_lang="id")).section_columns == 1


def test_render_emits_columns(tmp_path):
    from transdoc.ir import Block, BlockType, Confidence, Document
    from transdoc.regenerate.docx_out import render
    d = Document(source_path="x", mime="docx", section_columns=2)
    d.blocks = [Block(id="a", type=BlockType.PARAGRAPH, text="body", confidence=Confidence())]
    out = tmp_path / "o.docx"
    render(d, Config(target_lang="id"), str(out))
    assert _num_cols(docx.Document(str(out))) == 2
