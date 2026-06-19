# © 2026 Muhammad Mishbakhuz Zuhail. All rights reserved.
# Proprietary — source-available for reference only; no use, copying, or
# distribution without written permission. See LICENSE.
"""Paragraph spacing + indentation captured (docx) + rendered (docx/pdf)."""

from __future__ import annotations

import pytest

from transdoc.config import Config
from transdoc.ir import BBox, Block, BlockType, Confidence, Document, Style


def _doc():
    d = Document(source_path="x", mime="docx")
    d.blocks = [Block(id="1", type=BlockType.PARAGRAPH, text="indented spaced para",
                      bbox=BBox(x0=0, y0=0, x1=400, y1=20), confidence=Confidence(),
                      style=Style(space_before=12.0, space_after=6.0, indent_first=18.0,
                                  line_spacing=1.5))]
    return d


def test_docx_extract_captures_spacing(tmp_path):
    docx = pytest.importorskip("docx")
    from docx.shared import Pt
    from transdoc.extract.docx import extract
    dd = docx.Document()
    p = dd.add_paragraph("hello spaced")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.first_line_indent = Pt(20)
    path = tmp_path / "s.docx"
    dd.save(str(path))
    blk = next(b for b in extract(str(path), Config(target_lang="id")).blocks
               if b.type == BlockType.PARAGRAPH)
    assert round(blk.style.space_before) == 10
    assert round(blk.style.indent_first) == 20


def test_pdf_block_html_spacing():
    from transdoc.regenerate.pdf_out import _block_html
    html, _ = _block_html(_doc().blocks[0])
    assert "margin-top:12pt" in html and "text-indent:18pt" in html and "line-height:1.5" in html


def test_docx_render_spacing(tmp_path):
    docx = pytest.importorskip("docx")
    from transdoc.regenerate.docx_out import render
    out = tmp_path / "o.docx"
    render(_doc(), Config(target_lang="id"), str(out))
    p = next(p for p in docx.Document(str(out)).paragraphs if p.text.strip())
    assert round(p.paragraph_format.space_before.pt) == 12
