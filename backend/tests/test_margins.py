# SPDX-License-Identifier: AGPL-3.0-only
# Copyright (C) 2026 Muhammad Mishbakhuz Zuhail
"""Page margins captured (docx) + applied to docx output."""

from __future__ import annotations

import pytest


def test_docx_margins_roundtrip(tmp_path):
    docx = pytest.importorskip("docx")
    from docx.shared import Pt

    from transdoc.config import Config
    from transdoc.extract.docx import extract
    from transdoc.regenerate.docx_out import render

    dd = docx.Document()
    sec = dd.sections[0]
    sec.left_margin = Pt(90)
    sec.right_margin = Pt(54)
    dd.add_paragraph("body")
    src = tmp_path / "s.docx"
    dd.save(str(src))

    doc = extract(str(src), Config(target_lang="id"))
    assert round(doc.page_margins["left"]) == 90 and round(doc.page_margins["right"]) == 54

    out = tmp_path / "o.docx"
    render(doc, Config(target_lang="id"), str(out))
    s2 = docx.Document(str(out)).sections[0]
    assert round(s2.left_margin.pt) == 90 and round(s2.right_margin.pt) == 54
