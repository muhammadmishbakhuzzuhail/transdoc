# © 2026 Muhammad Mishbakhuz Zuhail. All rights reserved.
# Proprietary — source-available for reference only; no use, copying, or
# distribution without written permission. See LICENSE.
"""Tab stops + drop-cap (DOCX): captured and reproduced."""

from __future__ import annotations

import pytest

from transdoc.config import Config

docx = pytest.importorskip("docx")


def test_capture_tab_stops(tmp_path):
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Pt

    from transdoc.extract.docx import extract
    dd = docx.Document()
    p = dd.add_paragraph("name\tvalue")
    p.paragraph_format.tab_stops.add_tab_stop(Pt(200), WD_TAB_ALIGNMENT.RIGHT)
    f = tmp_path / "in.docx"
    dd.save(str(f))
    blk = next(b for b in extract(str(f), Config(target_lang="id")).blocks if "name" in b.text)
    assert blk.style.tab_stops
    pos, align = blk.style.tab_stops[0]
    assert round(pos) == 200 and align == "right"


def test_render_tab_stops(tmp_path):
    from transdoc.ir import Block, BlockType, Confidence, Document, Style
    from transdoc.regenerate.docx_out import render
    d = Document(source_path="x", mime="docx")
    d.blocks = [Block(id="a", type=BlockType.PARAGRAPH, text="x\ty",
                      style=Style(tab_stops=[(150.0, "right")]), confidence=Confidence())]
    out = tmp_path / "o.docx"
    render(d, Config(target_lang="id"), str(out))
    p = next(p for p in docx.Document(str(out)).paragraphs if p.text)
    assert len(p.paragraph_format.tab_stops._pPr.tabs) >= 1


def test_capture_render_drop_cap(tmp_path):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from transdoc.extract.docx import extract
    from transdoc.regenerate.docx_out import render
    dd = docx.Document()
    p = dd.add_paragraph("Once upon a time")
    ppr = p._p.get_or_add_pPr()
    fp = OxmlElement("w:framePr")
    fp.set(qn("w:dropCap"), "drop")
    fp.set(qn("w:lines"), "3")
    ppr.append(fp)
    f = tmp_path / "in.docx"
    dd.save(str(f))
    blk = next(b for b in extract(str(f), Config(target_lang="id")).blocks if "Once" in b.text)
    assert blk.style.drop_cap is True
    out = tmp_path / "o.docx"
    blk.translated = "Pada suatu waktu"
    from transdoc.ir import Document
    d = Document(source_path="x", mime="docx")
    d.blocks = [blk]
    render(d, Config(target_lang="id"), str(out))
    body = docx.Document(str(out)).element.body
    fps = body.findall(".//" + qn("w:framePr"))
    assert any(fp.get(qn("w:dropCap")) == "drop" for fp in fps)
