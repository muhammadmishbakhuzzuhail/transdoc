"""DOCX flow renderer carries IR character style (font/size/bold/italic/colour) + alignment —
previously dropped (plain 12pt black)."""

from __future__ import annotations

import pytest

pytest.importorskip("docx")

from docx import Document as Docx  # noqa: E402
from docx.shared import RGBColor  # noqa: E402

from transdoc.config import Config  # noqa: E402
from transdoc.ir import BBox, Block, BlockType, Confidence, Document, Style  # noqa: E402
from transdoc.regenerate.docx_out import render  # noqa: E402


def _doc():
    d = Document(source_path="x", mime="application/pdf")
    bb = BBox(x0=0, y0=0, x1=1, y1=1)
    d.blocks = [
        Block(id="1", type=BlockType.PARAGRAPH, text="bold blue centered", bbox=bb,
              confidence=Confidence(),
              style=Style(bold=True, color="#0000ff", size=15.0, align="center")),
        Block(id="2", type=BlockType.PARAGRAPH, text="italic justified", bbox=bb,
              confidence=Confidence(), style=Style(italic=True, align="justify")),
    ]
    return d


def test_docx_out_applies_character_style_and_alignment(tmp_path):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    out = tmp_path / "o.docx"
    render(_doc(), Config(target_lang="id"), str(out))
    paras = [p for p in Docx(str(out)).paragraphs if p.text.strip()]
    p1 = next(p for p in paras if "bold blue" in p.text)
    r1 = p1.runs[0]
    assert r1.font.bold is True
    assert r1.font.color.rgb == RGBColor(0x00, 0x00, 0xff)
    assert round(r1.font.size.pt) == 15
    assert p1.alignment == WD_ALIGN_PARAGRAPH.CENTER

    p2 = next(p for p in paras if "italic justified" in p.text)
    assert p2.runs[0].font.italic is True
    assert p2.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY


def test_docx_out_renders_merged_cells(tmp_path):
    from transdoc.ir import Cell, Table
    d = Document(source_path="x", mime="application/pdf")
    bb = BBox(x0=0, y0=0, x1=1, y1=1)
    # row0: a header spanning 2 cols; row1: two cells
    tbl = Table(rows=[[Cell(text="Header", colspan=2)],
                      [Cell(text="a"), Cell(text="b")]])
    d.blocks = [Block(id="t", type=BlockType.TABLE, bbox=bb, table=tbl, confidence=Confidence())]
    out = tmp_path / "m.docx"
    render(d, Config(target_lang="id"), str(out))
    t = Docx(str(out)).tables[0]
    assert len(t.columns) == 2 and len(t.rows) == 2
    # merged header cell spans both columns -> both grid cells resolve to same text
    assert t.cell(0, 0).text == "Header" and t.cell(0, 1).text == "Header"
    assert t.cell(1, 0).text == "a" and t.cell(1, 1).text == "b"
