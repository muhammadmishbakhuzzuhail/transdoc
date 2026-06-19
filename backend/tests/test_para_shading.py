# SPDX-License-Identifier: AGPL-3.0-only
# Copyright (C) 2026 Muhammad Mishbakhuz Zuhail
"""Paragraph shading + border (callout/boxed paragraphs): captured from DOCX, rendered to
both DOCX and PDF-HTML."""

from __future__ import annotations

import pytest

from transdoc.config import Config


def test_capture_shading_and_border(tmp_path):
    docx = pytest.importorskip("docx")
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    from transdoc.extract.docx import extract
    dd = docx.Document()
    p = dd.add_paragraph("callout paragraph")
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "FFFF00")
    ppr.append(shd)
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        pbdr.append(e)
    ppr.append(pbdr)
    f = tmp_path / "in.docx"
    dd.save(str(f))
    blk = next(b for b in extract(str(f), Config(target_lang="id")).blocks
               if b.text == "callout paragraph")
    assert blk.style.para_shading == "#FFFF00"
    assert blk.style.para_border is True


def _styled_doc():
    from transdoc.ir import BBox, Block, BlockType, Confidence, Document, Style
    d = Document(source_path="x", mime="docx")
    d.blocks = [Block(id="a", type=BlockType.PARAGRAPH, text="boxed",
                      bbox=BBox(x0=0, y0=0, x1=1, y1=1),
                      style=Style(para_shading="#FFFF00", para_border=True),
                      confidence=Confidence())]
    return d


def test_docx_renders_shading_border(tmp_path):
    docx = pytest.importorskip("docx")
    from docx.oxml.ns import qn

    from transdoc.regenerate.docx_out import render
    out = tmp_path / "o.docx"
    render(_styled_doc(), Config(target_lang="id"), str(out))
    body = docx.Document(str(out)).element.body
    assert body.findall(".//" + qn("w:shd")) and body.findall(".//" + qn("w:pBdr"))


def test_pdf_html_has_shading_border():
    from transdoc.regenerate.pdf_out import _block_html
    html, _ = _block_html(_styled_doc().blocks[0])
    assert "background-color:#FFFF00" in html and "border:1px solid #000" in html
