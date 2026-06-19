# © 2026 Muhammad Mishbakhuz Zuhail. All rights reserved.
# Proprietary — source-available for reference only; no use, copying, or
# distribution without written permission. See LICENSE.
"""RTL / bidi support: detect direction from the TRANSLATED text + target language, propagate
Style.rtl, and emit the right direction markers per renderer (DOCX w:bidi/w:rtl, EPUB dir,
PDF raw-draw reshaping).

The core bug this guards: an LTR source translated INTO an RTL language (EN->AR) used to render
flush-left with no base direction, because nothing recomputed direction from the output.
"""

from __future__ import annotations

import pytest

from transdoc import textdir
from transdoc.config import Config
from transdoc.ir import Block, BlockType, Cell, Document, Style, Table

AR = "مرحبا بالعالم"          # pure Arabic
HE = "שלום עולם"              # pure Hebrew
EN = "Hello world"


# ---- detection ----------------------------------------------------------------------------

def test_is_rtl_lang():
    for code in ("ar", "he", "fa", "ur", "ps", "ckb", "ar-EG", "fa_IR", "HE"):
        assert textdir.is_rtl_lang(code), code
    for code in ("en", "id", "fr", "zh", "ja", "ru", None, ""):
        assert not textdir.is_rtl_lang(code), code


def test_rtl_ratio_and_text():
    assert textdir.rtl_ratio(AR) == 1.0
    assert textdir.rtl_ratio(EN) == 0.0
    assert textdir.rtl_ratio("123 ... ") == 0.0      # neutrals don't count
    assert textdir.is_rtl_text(AR)
    assert not textdir.is_rtl_text(EN)
    assert not textdir.is_rtl_text("Hello world example مر")   # Latin-dominant -> not RTL


def test_is_mixed_bidi():
    assert textdir.is_mixed_bidi("مرحبا Google")
    assert not textdir.is_mixed_bidi(AR)
    assert not textdir.is_mixed_bidi(EN)
    assert not textdir.is_mixed_bidi("مرحبا 123")     # digits are neutral, not strong-LTR


def test_effective_rtl_text_wins_over_lang():
    # text with its own direction overrides the target language
    assert textdir.effective_rtl(AR, "en") is True
    assert textdir.effective_rtl(EN, "ar") is False    # a Latin URL stays LTR in an Arabic doc
    # empty / non-directional text falls back to the target language
    assert textdir.effective_rtl("", "ar") is True
    assert textdir.effective_rtl("123", "he") is True
    assert textdir.effective_rtl("", "en") is False


# ---- shaping (raw-draw paths) -------------------------------------------------------------

def test_shape_for_raw_draw():
    assert textdir.shape_for_raw_draw(EN, rtl=False) == EN     # no-op for LTR
    shaped = textdir.shape_for_raw_draw(AR, rtl=True)
    pytest.importorskip("bidi")
    pytest.importorskip("arabic_reshaper")
    # reshape+reorder changes the codepoint sequence (presentation forms + visual order)
    assert shaped != AR
    assert len(shaped) >= 1


# ---- propagation --------------------------------------------------------------------------

def _doc(*blocks: Block) -> Document:
    d = Document(source_path="x.pdf", mime="application/pdf")
    d.blocks = list(blocks)
    return d


def test_apply_text_direction_sets_rtl_from_output():
    ar = Block(id="a", type=BlockType.PARAGRAPH, text="Hello", translated=AR, style=Style())
    en = Block(id="b", type=BlockType.PARAGRAPH, text="Bonjour", translated=EN, style=Style())
    doc = _doc(ar, en)
    textdir.apply_text_direction(doc, Config(target_lang="ar"))
    assert ar.style.rtl is True
    assert en.style.rtl is False                # Latin output stays LTR even with ar target


def test_apply_text_direction_runs_and_cells():
    from transdoc.ir import Run
    b = Block(id="t", type=BlockType.TABLE, style=Style())
    b.runs = [Run(text="x", translated=AR, style=Style()),
              Run(text="y", translated=EN, style=Style())]
    b.table = Table(rows=[[Cell(text="h", translated=AR), Cell(text="k", translated=EN)]])
    textdir.apply_text_direction(_doc(b), Config(target_lang="ar"))
    assert b.runs[0].style.rtl is True
    assert b.runs[1].style.rtl is False
    assert b.table.rows[0][0].align == "right"   # RTL cell right-aligned
    assert b.table.rows[0][1].align is None       # LTR cell untouched


# ---- DOCX renderer ------------------------------------------------------------------------

def test_docx_emits_bidi_and_rtl(tmp_path):
    pytest.importorskip("docx")
    from docx import Document as Docx
    from docx.oxml.ns import qn

    from transdoc.regenerate.docx_out import render

    b = Block(id="a", type=BlockType.PARAGRAPH, text="Hello", translated=AR,
              style=Style(rtl=True))
    out = tmp_path / "out.docx"
    render(_doc(b), Config(target_lang="ar"), str(out))

    d = Docx(str(out))
    paras = [p for p in d.paragraphs if p.text.strip()]
    assert paras, "no body paragraph emitted"
    p = paras[0]
    assert p._p.find(qn("w:pPr")).find(qn("w:bidi")) is not None       # paragraph bidi
    assert any(r._r.find(qn("w:rPr")) is not None
               and r._r.find(qn("w:rPr")).find(qn("w:rtl")) is not None
               for r in p.runs)                                         # run rtl


def test_docx_ltr_has_no_bidi(tmp_path):
    pytest.importorskip("docx")
    from docx import Document as Docx
    from docx.oxml.ns import qn

    from transdoc.regenerate.docx_out import render

    b = Block(id="a", type=BlockType.PARAGRAPH, text="Hi", translated=EN, style=Style(rtl=False))
    out = tmp_path / "ltr.docx"
    render(_doc(b), Config(target_lang="en"), str(out))
    d = Docx(str(out))
    p = [p for p in d.paragraphs if p.text.strip()][0]
    ppr = p._p.find(qn("w:pPr"))
    assert ppr is None or ppr.find(qn("w:bidi")) is None
