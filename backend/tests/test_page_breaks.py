# © 2026 Muhammad Mishbakhuz Zuhail. All rights reserved.
# Proprietary — source-available for reference only; no use, copying, or
# distribution without written permission. See LICENSE.
"""Manual page breaks (DOCX): captured from the source and reproduced in the output."""

from __future__ import annotations

import pytest

from transdoc.config import Config

docx = pytest.importorskip("docx")


def _has_page_break(document) -> bool:
    from docx.oxml.ns import qn
    for br in document.element.body.findall(".//" + qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def test_capture_pagebreakbefore(tmp_path):
    from transdoc.extract.docx import extract
    dd = docx.Document()
    dd.add_paragraph("first page")
    p2 = dd.add_paragraph("second page")
    p2.paragraph_format.page_break_before = True
    f = tmp_path / "in.docx"
    dd.save(str(f))
    blocks = extract(str(f), Config(target_lang="id")).blocks
    by_text = {b.text: b for b in blocks}
    assert by_text["second page"].page_break_before is True
    assert by_text["first page"].page_break_before is False


def test_capture_run_page_break(tmp_path):
    from transdoc.extract.docx import extract
    dd = docx.Document()
    p = dd.add_paragraph("before")
    p.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
    f = tmp_path / "in2.docx"
    dd.save(str(f))
    blk = next(b for b in extract(str(f), Config(target_lang="id")).blocks if b.text == "before")
    assert blk.page_break_before is True


def test_render_reproduces_break(tmp_path):
    from transdoc.ir import BBox, Block, BlockType, Confidence, Document
    from transdoc.regenerate.docx_out import render
    d = Document(source_path="x", mime="docx")
    d.blocks = [
        Block(id="a", type=BlockType.PARAGRAPH, text="one", bbox=BBox(x0=0, y0=0, x1=1, y1=1),
              confidence=Confidence()),
        Block(id="b", type=BlockType.PARAGRAPH, text="two", page_break_before=True,
              bbox=BBox(x0=0, y0=2, x1=1, y1=3), confidence=Confidence()),
    ]
    out = tmp_path / "o.docx"
    render(d, Config(target_lang="id"), str(out))
    assert _has_page_break(docx.Document(str(out)))
