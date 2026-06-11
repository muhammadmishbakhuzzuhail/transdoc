"""Bilingual output: source + translation together, across renderers (was markdown-only)."""

from __future__ import annotations

import fitz

from transdoc.config import Config
from transdoc.ir import Block, BlockType, Confidence, Document
from transdoc.regenerate.docx_out import render as docx_render
from transdoc.regenerate.markdown import render as md_render
from transdoc.regenerate.pdf_out import render_flow


def _doc() -> Document:
    d = Document(source_path="x.docx", mime="docx", target_lang="id")
    d.blocks = [Block(id="b0", type=BlockType.PARAGRAPH, page=0,
                      text="Hello world", translated="Halo dunia",
                      confidence=Confidence(source="digital"))]
    return d


def test_markdown_bilingual_has_both():
    out = md_render(_doc(), Config(target_lang="id", bilingual=True))
    assert "Hello world" in out and "Halo dunia" in out


def test_docx_bilingual_has_both(tmp_path):
    out = tmp_path / "o.docx"
    docx_render(_doc(), Config(target_lang="id", bilingual=True), str(out))
    import docx
    paras = [p.text for p in docx.Document(str(out)).paragraphs]
    assert "Hello world" in paras       # source line, untranslated
    assert "Halo dunia" in paras        # translation line


def test_docx_non_bilingual_omits_source(tmp_path):
    out = tmp_path / "o.docx"
    docx_render(_doc(), Config(target_lang="id", bilingual=False), str(out))
    import docx
    text = "\n".join(p.text for p in docx.Document(str(out)).paragraphs)
    assert "Halo dunia" in text
    assert "Hello world" not in text    # only the translation


def test_pdf_flow_bilingual_has_both(tmp_path):
    out = tmp_path / "o.pdf"
    render_flow(_doc(), Config(target_lang="id", bilingual=True), str(out))
    text = "".join(p.get_text() for p in fitz.open(str(out)))
    assert "Hello world" in text and "Halo dunia" in text
