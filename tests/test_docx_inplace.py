"""In-place DOCX translation (DeepL strategy): a .docx source is mutated, swapping text while
keeping every paragraph style, list, and table — not rebuilt from the IR."""

from __future__ import annotations

import pytest

pytest.importorskip("docx")

from docx import Document as Docx  # noqa: E402

from transdoc.config import Config, Engine, Mode, OutputFormat  # noqa: E402
from transdoc.pipeline import run  # noqa: E402


def _make_docx(path):
    d = Docx()
    d.add_heading("Project Report", level=1)
    d.add_paragraph("This is the body paragraph with some content.")
    d.add_paragraph("First item", style="List Bullet")
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Name"
    t.rows[0].cells[1].text = "Value"
    t.rows[1].cells[0].text = "Alpha"
    t.rows[1].cells[1].text = "One"
    d.save(str(path))


def test_inplace_keeps_styles_and_swaps_text(tmp_path):
    src = tmp_path / "doc.docx"
    _make_docx(src)
    out = tmp_path / "doc.id.docx"
    run(str(src), Config(source_lang="en", target_lang="id", engine=Engine.ECHO,
                         output_format=OutputFormat.DOCX, mode=Mode.FULL), out_path=str(out))

    o, t = Docx(str(src)), Docx(str(out))
    # structure preserved 1:1
    assert len(o.paragraphs) == len(t.paragraphs)
    assert len(o.tables) == len(t.tables)
    # styles preserved, text translated (echo prefixes "[id] ")
    by_style = {p.style.name: p.text for p in t.paragraphs if p.text.strip()}
    assert any("Heading" in s and v.startswith("[id]") for s, v in by_style.items())
    assert any("List" in s for s in by_style)
    # table cell translated in place
    assert t.tables[0].rows[0].cells[0].text == "[id] Name"
